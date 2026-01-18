from __future__ import annotations
import re
import base64
import datetime
import time
import io
import asyncio
import logging
import hashlib
import struct
import zlib
import binascii
from pathlib import Path
from dataclasses import dataclass
from typing import Optional, Callable, Awaitable, Any, List, Tuple, Dict, cast
from urllib.parse import quote
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import parse_xml
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from open_webui.models.chats import Chats
from open_webui.models.users import Users
from open_webui.utils.chat import generate_chat_completion
from pydantic import BaseModel, Field
# Files are used to embed internal /api/v1/files/<id>/content images.
try:
    from open_webui.models.files import Files  # type: ignore
except Exception:  # pragma: no cover - depends on host Open WebUI runtime
    Files = None
# Pygments for syntax highlighting
try:
    from pygments import lex
    from pygments.lexers import get_lexer_by_name, TextLexer
    from pygments.token import Token
    PYGMENTS_AVAILABLE = True
except ImportError:
    PYGMENTS_AVAILABLE = False
try:
    from latex2mathml.converter import convert as latex_to_mathml
    import mathml2omml
    LATEX_MATH_AVAILABLE = True
except Exception:
    LATEX_MATH_AVAILABLE = False
# boto3 for S3 direct access (faster than API fallback)
try:
    import boto3
    from botocore.config import Config as BotoConfig
    import os
    BOTO3_AVAILABLE = True
except ImportError:
    BOTO3_AVAILABLE = False

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
)
logger = logging.getLogger(__name__)
_AUTO_URL_RE = re.compile(r"(?:https?://|www\.)[^\s<>()]+")
_DATA_IMAGE_URL_RE = re.compile(
    r"^data:(?P<mime>image/[a-z0-9.+-]+)\s*;\s*base64\s*,\s*(?P<b64>.*)$",
    re.IGNORECASE | re.DOTALL,
)
_OWUI_API_FILE_ID_RE = re.compile(
    r"/api/v1/files/(?P<id>[A-Za-z0-9-]+)(?:/content)?(?:[/?#]|$)",
    re.IGNORECASE,
)
_CURRENCY_NUMBER_RE = re.compile(r"^\d[\d,]*(?:\.\d+)?$")
_TRANSPARENT_1PX_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVQImWNgYGBgAAAABQABDQottAAAAABJRU5ErkJggg=="
)
_ASVG_NS = "http://schemas.microsoft.com/office/drawing/2016/SVG/main"
nsmap.setdefault("asvg", _ASVG_NS)
_REASONING_DETAILS_RE = re.compile(
    r"<details\b[^>]*\btype\s*=\s*(?:\"reasoning\"|'reasoning'|reasoning)[^>]*>.*?</details\s*>",
    re.IGNORECASE | re.DOTALL,
)
_THINK_RE = re.compile(r"<think\b[^>]*>.*?</think\s*>", re.IGNORECASE | re.DOTALL)
_ANALYSIS_RE = re.compile(
    r"<analysis\b[^>]*>.*?</analysis\s*>", re.IGNORECASE | re.DOTALL
)

@dataclass(frozen=True)
class _CitationRef:
    idx: int
    anchor: str
    title: str
    url: Optional[str]
    source_id: str

class Action:
    # Internationalization message dictionaries
    _I18N_MESSAGES: Dict[str, Dict[str, str]] = {
        "en": {
            "converting": "Converting to Word document...",
            "exported": "Word document exported",
            "success": "Successfully exported to {filename}",
            "error_no_content": "No content found to export!",
            "error_export": "Error exporting Word document: {error}",
            "export_failed": "Export failed: {error}",
            "figure_prefix": "Figure",
            "references": "References",
        },
        "zh": {
            "converting": "正在转换为 Word 文档...",
            "exported": "Word 文档导出完成",
            "success": "成功导出至 {filename}",
            "error_no_content": "没有找到可导出的内容！",
            "error_export": "导出 Word 文档时出错: {error}",
            "export_failed": "导出失败: {error}",
            "figure_prefix": "图",
            "references": "参考文献",
        },
    }
    class Valves(BaseModel):
        TITLE_SOURCE: str = Field(
            default="chat_title",
            description="Title Source: 'chat_title' (Chat Title), 'ai_generated' (AI Generated), 'markdown_title' (Markdown Title)",
        )
        SHOW_STATUS: bool = Field(
            default=True,
            description="Whether to show operation status updates.",
        )
        SHOW_DEBUG_LOG: bool = Field(
            default=False,
            description="Whether to print debug logs in the browser console.",
        )
        MAX_EMBED_IMAGE_MB: int = Field(
            default=20,
            description="Maximum image size to embed into DOCX (MB). Applies to data URLs and /api/v1/files/<id>/content images.",
        )
        # Font configuration
        FONT_LATIN: str = Field(
            default="Times New Roman",
            description="Font for Latin characters (e.g., 'Times New Roman', 'Calibri', 'Arial')",
        )
        FONT_ASIAN: str = Field(
            default="SimSun",
            description="Font for Asian characters (e.g., 'SimSun', 'Microsoft YaHei', 'PingFang SC')",
        )
        FONT_CODE: str = Field(
            default="Consolas",
            description="Font for code blocks and inline code (e.g., 'Consolas', 'Courier New', 'Monaco')",
        )
        # Table styling
        TABLE_HEADER_COLOR: str = Field(
            default="F2F2F2",
            description="Table header background color (hex, without #)",
        )
        TABLE_ZEBRA_COLOR: str = Field(
            default="FBFBFB",
            description="Table zebra stripe background color for alternate rows (hex, without #)",
        )
        MERMAID_JS_URL: str = Field(
            default="https://cdn.jsdelivr.net/npm/mermaid@11.12.2/dist/mermaid.min.js",
            description="Mermaid JS CDN URL",
        )
        MERMAID_JSZIP_URL: str = Field(
            default="https://cdnjs.cloudflare.com/ajax/libs/jszip/3.10.1/jszip.min.js",
            description="JSZip CDN URL (DOCX manipulation)",
        )
        MERMAID_PNG_SCALE: float = Field(
            default=3.0,
            description="PNG render resolution multiplier (higher = clearer, larger file)",
        )
        MERMAID_DISPLAY_SCALE: float = Field(
            default=1.0,
            description="Diagram width relative to available page width (<=1 recommended)",
        )
        MERMAID_OPTIMIZE_LAYOUT: bool = Field(
            default=False,
            description="Optimize Mermaid layout: convert LR to TD for graph/flowchart",
        )
        MERMAID_BACKGROUND: str = Field(
            default="",
            description="Mermaid background color. Empty = transparent (recommended for Word dark mode). Used only for optional PNG fill.",
        )
        MERMAID_CAPTIONS_ENABLE: bool = Field(
            default=True,
            description="Add figure captions under Mermaid images/charts",
        )
        MERMAID_CAPTION_STYLE: str = Field(
            default="Caption",
            description="Paragraph style name for Mermaid captions (uses 'Caption' if available, otherwise creates a safe custom style)",
        )
        MERMAID_CAPTION_PREFIX: str = Field(
            default="",
            description="Caption prefix label (e.g., 'Figure' or '图'). Empty = auto-detect based on user language.",
        )
        MATH_ENABLE: bool = Field(
            default=True,
            description="Enable LaTeX math block conversion (\\[...\\] and $$...$$) into Word equations",
        )
        MATH_INLINE_DOLLAR_ENABLE: bool = Field(
            default=True,
            description="Enable inline $...$ math conversion into Word equations (conservative parsing to reduce false positives)",
        )
        # Language configuration
        UI_LANGUAGE: str = Field(
            default="en",
            description="UI language for export messages. Options: 'en' (English), 'zh' (Chinese)",
        )
    class UserValves(BaseModel):
        TITLE_SOURCE: str = Field(
            default="chat_title",
            description="Title Source: 'chat_title' (Chat Title), 'ai_generated' (AI Generated), 'markdown_title' (Markdown Title)",
        )
        UI_LANGUAGE: str = Field(
            default="en",
            description="UI language for export messages. Options: 'en' (English), 'zh' (Chinese)",
        )
        FONT_LATIN: str = Field(
            default="Times New Roman",
            description="Font for Latin characters (e.g., 'Times New Roman', 'Calibri', 'Arial')",
        )
        FONT_ASIAN: str = Field(
            default="SimSun",
            description="Font for Asian characters (e.g., 'SimSun', 'Microsoft YaHei', 'PingFang SC')",
        )
        FONT_CODE: str = Field(
            default="Consolas",
            description="Font for code blocks and inline code (e.g., 'Consolas', 'Courier New', 'Monaco')",
        )
        TABLE_HEADER_COLOR: str = Field(
            default="F2F2F2",
            description="Table header background color (hex, without #)",
        )
        TABLE_ZEBRA_COLOR: str = Field(
            default="FBFBFB",
            description="Table zebra stripe background color for alternate rows (hex, without #)",
        )
        MERMAID_PNG_SCALE: float = Field(
            default=3.0,
            description="PNG render resolution multiplier (higher = clearer, larger file)",
        )
        MERMAID_DISPLAY_SCALE: float = Field(
            default=1.0,
            description="Diagram width relative to available page width (<=1 recommended)",
        )
        MERMAID_OPTIMIZE_LAYOUT: bool = Field(
            default=False,
            description="Optimize Mermaid layout: convert LR to TD for graph/flowchart",
        )
        MERMAID_BACKGROUND: str = Field(
            default="",
            description="Mermaid background color. Empty = transparent (recommended for Word dark mode). Used only for optional PNG fill.",
        )
        MERMAID_CAPTIONS_ENABLE: bool = Field(
            default=True,
            description="Add figure captions under Mermaid images/charts",
        )
        MATH_ENABLE: bool = Field(
            default=True,
            description="Enable LaTeX math block conversion (\\\\[...\\\\] and $$...$$) into Word equations",
        )
        MATH_INLINE_DOLLAR_ENABLE: bool = Field(
            default=True,
            description="Enable inline $...$ math conversion into Word equations (conservative parsing to reduce false positives)",
        )
    def __init__(self):
        self.valves = self.Valves()
        self._mermaid_figure_counter: int = 0
        self._mermaid_placeholder_counter: int = 0
        self._caption_style_name: Optional[str] = None
        self._citation_anchor_by_index: Dict[int, str] = {}
        self._citation_refs: List[_CitationRef] = []
        self._bookmark_id_counter: int = 1
        self._active_doc: Optional[Document] = None
        self._user_lang: str = "en"  # Will be set per-request
        self._api_token: Optional[str] = None
        self._api_base_url: Optional[str] = None
    def _get_lang_key(self, user_language: str) -> str:
        """Convert user language code to i18n key (e.g., 'zh-CN' -> 'zh', 'en-US' -> 'en')."""
        lang = (user_language or "en").lower().split("-")[0]
        return lang if lang in self._I18N_MESSAGES else "en"
    def _get_msg(self, key: str, **kwargs) -> str:
        """Get internationalized message by key with optional formatting."""
        messages = self._I18N_MESSAGES.get(self._user_lang, self._I18N_MESSAGES["en"])
        msg = messages.get(key, self._I18N_MESSAGES["en"].get(key, key))
        if kwargs:
            try:
                return msg.format(**kwargs)
            except KeyError:
                return msg
        return msg
    def _get_user_context(self, __user__: Optional[Dict[str, Any]]) -> Dict[str, str]:
        """Safely extracts user context information."""
        if isinstance(__user__, (list, tuple)):
            user_data = __user__[0] if __user__ else {}
        elif isinstance(__user__, dict):
            user_data = __user__
        else:
            user_data = {}
        return {
            "user_id": user_data.get("id", "unknown_user"),
            "user_name": user_data.get("name", "User"),
            "user_language": user_data.get("language", "en-US"),
        }
    def _get_chat_context(
        self, body: dict, __metadata__: Optional[dict] = None
    ) -> Dict[str, str]:
        """
        Unified extraction of chat context information (chat_id, message_id).
        Prioritizes extraction from body, then metadata.
        """
        chat_id = ""
        message_id = ""
        # 1. Try to get from body
        if isinstance(body, dict):
            chat_id = body.get("chat_id", "")
            message_id = body.get("id", "")  # message_id is usually 'id' in body
            # Check body.metadata as fallback
            if not chat_id or not message_id:
                body_metadata = body.get("metadata", {})
                if isinstance(body_metadata, dict):
                    if not chat_id:
                        chat_id = body_metadata.get("chat_id", "")
                    if not message_id:
                        message_id = body_metadata.get("message_id", "")
        # 2. Try to get from __metadata__ (as supplement)
        if __metadata__ and isinstance(__metadata__, dict):
            if not chat_id:
                chat_id = __metadata__.get("chat_id", "")
            if not message_id:
                message_id = __metadata__.get("message_id", "")
        return {
            "chat_id": str(chat_id).strip(),
            "message_id": str(message_id).strip(),
        }
    async def _emit_status(
        self,
        emitter: Optional[Callable[[Any], Awaitable[None]]],
        description: str,
        done: bool = False,
    ):
        """Emits a status update event."""
        if self.valves.SHOW_STATUS and emitter:
            await emitter(
                {"type": "status", "data": {"description": description, "done": done}}
            )
    async def _emit_notification(
        self,
        emitter: Optional[Callable[[Any], Awaitable[None]]],
        content: str,
        ntype: str = "info",
    ):
        """Emits a notification event (info, success, warning, error)."""
        if emitter:
            await emitter(
                {"type": "notification", "data": {"type": ntype, "content": content}}
            )
    async def _emit_debug_log(self, emitter, title: str, data: dict):
        """Print structured debug logs in the browser console"""
        if not self.valves.SHOW_DEBUG_LOG or not emitter:
            return
        try:
            import json
            js_code = f"""
                (async function() {{
                    console.group("🛠️ {title}");
                    console.log({json.dumps(data, ensure_ascii=False)});
                    console.groupEnd();
                }})();
            """
            await emitter({"type": "execute", "data": {"code": js_code}})
        except Exception as e:
            print(f"Error emitting debug log: {e}")
    async def action(
        self,
        body: dict,
        __user__=None,
        __event_emitter__=None,
        __event_call__: Optional[Callable[[Any], Awaitable[None]]] = None,
        __metadata__: Optional[dict] = None,
        __request__: Optional[Any] = None,
    ):
        logger.info(f"action:{__name__}")
        # Parse user info
        user_name = "User"
        user_id = "unknown_user"
        if isinstance(__user__, (list, tuple)):
            user_name = __user__[0].get("name", "User") if __user__[0] else "User"
            user_id = (
                __user__[0]["id"]
                if __user__ and "id" in __user__[0]
                else "unknown_user"
            )
        elif isinstance(__user__, dict):
            user_name = __user__.get("name", "User")
            user_id = __user__.get("id", "unknown_user")
        # Apply UserValves if present
        if __user__ and "valves" in __user__:
            # Update self.valves with user-specific values
            # Note: This assumes per-request instantiation or that we are okay with modifying the singleton.
            # Given the plugin architecture, we'll update it for this execution.
            for key, value in __user__["valves"].model_dump().items():
                if hasattr(self.valves, key):
                    setattr(self.valves, key, value)
        # Get user language from Valves configuration
        self._user_lang = self._get_lang_key(self.valves.UI_LANGUAGE)
        # Extract API connection info for file fetching (S3/Object Storage support)
        def _get_default_base_url() -> str:
            port = os.environ.get("PORT") or "8080"
            return f"http://localhost:{port}"
        if __request__:
            try:
                self._api_token = __request__.headers.get("Authorization")
                self._api_base_url = str(__request__.base_url).rstrip("/")
            except Exception:
                self._api_token = None
                self._api_base_url = _get_default_base_url()
        else:
            self._api_token = None
            self._api_base_url = _get_default_base_url()
        if __event_emitter__:
            last_assistant_message = body["messages"][-1]
            await __event_emitter__(
                {
                    "type": "status",
                    "data": {
                        "description": self._get_msg("converting"),
                        "done": False,
                    },
                }
            )
            try:
                message_content = last_assistant_message["content"]
                if isinstance(message_content, str):
                    message_content = self._strip_reasoning_blocks(message_content)
                if not message_content or not message_content.strip():
                    await self._emit_notification(
                        __event_emitter__, self._get_msg("error_no_content"), "error"
                    )
                    return
                # Generate filename
                title = ""
                chat_ctx = self._get_chat_context(body, __metadata__)
                chat_id = chat_ctx["chat_id"]
                # Fetch chat_title directly via chat_id as it's usually missing in body
                chat_title = ""
                if chat_id:
                    chat_title = await self.fetch_chat_title(chat_id, user_id)
                if (
                    self.valves.TITLE_SOURCE.strip() == "chat_title"
                    or not self.valves.TITLE_SOURCE.strip()
                ):
                    title = chat_title
                elif self.valves.TITLE_SOURCE.strip() == "markdown_title":
                    title = self.extract_title(message_content)
                elif self.valves.TITLE_SOURCE.strip() == "ai_generated":
                    title = await self.generate_title_using_ai(
                        body, message_content, user_id, __request__
                    )
                # Fallback logic
                if not title:
                    if self.valves.TITLE_SOURCE.strip() != "chat_title" and chat_title:
                        title = chat_title
                    elif self.valves.TITLE_SOURCE.strip() != "markdown_title":
                        extracted = self.extract_title(message_content)
                        if extracted:
                            title = extracted
                current_datetime = datetime.datetime.now()
                formatted_date = current_datetime.strftime("%Y%m%d")
                cleaned_title = self.clean_filename(title) if title else ""
                if cleaned_title:
                    filename = f"{cleaned_title}.docx"
                else:
                    clean_user = self.clean_filename(user_name)
                    filename = f"{clean_user}_{formatted_date}.docx"
                # Escape filename for JS string
                js_filename = filename.replace("\\", "\\\\").replace('"', '\\"')
                top_heading = ""
                if chat_title:
                    top_heading = chat_title
                elif title:
                    top_heading = title
                # Create Word document; if no h1 exists, inject chat title as h1
                has_h1 = bool(re.search(r"^#\s+.+$", message_content, re.MULTILINE))
                sources = (
                    last_assistant_message.get("sources") or body.get("sources") or []
                )
                doc = await self.markdown_to_docx(
                    message_content,
                    top_heading=top_heading,
                    has_h1=has_h1,
                    sources=sources,
                    event_emitter=__event_emitter__,
                )
                # Save to memory
                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)
                file_content = doc_buffer.read()
                base64_blob = base64.b64encode(file_content).decode("utf-8")
                # Trigger file download
                if __event_call__:
                    await __event_call__(
                        {
                            "type": "execute",
                            "data": {
                                "code": f"""
                                (async function() {{
                                    const base64Data = "{base64_blob}";
                                    const filename = "{js_filename}";
	                                    const mermaidUrl = "{self.valves.MERMAID_JS_URL}";
	                                    const jszipUrl = "{self.valves.MERMAID_JSZIP_URL}";
	                                    const pngScale = {float(self.valves.MERMAID_PNG_SCALE)};
	                                    const displayScale = {float(self.valves.MERMAID_DISPLAY_SCALE)};
	                                    const bgRaw = "{(self.valves.MERMAID_BACKGROUND or '').strip()}";
	                                    const bg = (bgRaw || "").trim();
	                                    const bgFill = (bg && bg.toLowerCase() !== "transparent") ? bg : "";
	                                    const themeBackground = bgFill || "transparent";
                                    function downloadBlob(blob, filename) {{
                                        const url = URL.createObjectURL(blob);
                                        const a = document.createElement("a");
                                        a.style.display = "none";
                                        a.href = url;
                                        a.download = filename;
                                        document.body.appendChild(a);
                                        a.click();
                                        URL.revokeObjectURL(url);
                                        document.body.removeChild(a);
                                    }}
                                    async function loadScript(url, globalName) {{
                                        if (globalName && window[globalName]) return;
                                        await new Promise((resolve, reject) => {{
                                            const script = document.createElement("script");
                                            script.src = url;
                                            script.onload = resolve;
                                            script.onerror = reject;
                                            document.head.appendChild(script);
                                        }});
                                    }}
                                    function decodeBase64ToUint8Array(b64) {{
                                        const binary = atob(b64);
                                        const bytes = new Uint8Array(binary.length);
                                        for (let i = 0; i < binary.length; i++) bytes[i] = binary.charCodeAt(i);
                                        return bytes;
                                    }}
                                    function parseViewBox(vb) {{
                                        if (!vb) return null;
                                        const parts = vb.trim().split(/\\s+/).map(Number);
                                        if (parts.length !== 4 || parts.some((n) => !isFinite(n))) return null;
                                        return {{ minX: parts[0], minY: parts[1], width: parts[2], height: parts[3] }};
                                    }}
	                                    function normalizeSvgForWord(svgText) {{
	                                        const parser = new DOMParser();
	                                        const doc = parser.parseFromString(svgText, "image/svg+xml");
	                                        const svgEl = doc.documentElement;
	                                        if (!svgEl || svgEl.tagName.toLowerCase() !== "svg") return svgText;
                                        // Pad viewBox a little to reduce clipping in Word.
                                        const vb0 = parseViewBox(svgEl.getAttribute("viewBox"));
                                        if (vb0 && vb0.width > 0 && vb0.height > 0) {{
                                            const minDim = Math.min(vb0.width, vb0.height);
                                            let pad = Math.max(8.0, minDim * 0.02);
                                            pad = Math.min(pad, 24.0);
                                            const vb = {{
                                                minX: vb0.minX - pad,
                                                minY: vb0.minY - pad,
                                                width: vb0.width + 2 * pad,
                                                height: vb0.height + 2 * pad,
                                            }};
                                            svgEl.setAttribute("viewBox", `${{vb.minX}} ${{vb.minY}} ${{vb.width}} ${{vb.height}}`);
                                        }}
                                        const vb = parseViewBox(svgEl.getAttribute("viewBox"));
                                        const widthAttr = (svgEl.getAttribute("width") || "").trim();
                                        const heightAttr = (svgEl.getAttribute("height") || "").trim();
                                        const widthPct = widthAttr.endsWith("%");
                                        const heightPct = heightAttr.endsWith("%");
                                        if (vb && vb.width > 0 && vb.height > 0 && (!widthAttr || !heightAttr || widthPct || heightPct)) {{
                                            svgEl.setAttribute("width", `${{vb.width}}`);
                                            svgEl.setAttribute("height", `${{vb.height}}`);
                                        }}
	                                        svgEl.removeAttribute("style");
	                                        svgEl.setAttribute("preserveAspectRatio", "xMidYMid meet");
		                                        svgEl.setAttribute("overflow", "visible");
		                                        const removeNode = (n) => {{
		                                            try {{ n && n.parentNode && n.parentNode.removeChild(n); }} catch (_e) {{}}
		                                        }};
		                                        // Remove Mermaid/OWUI background rectangles to avoid \"white box\" rendering in Word dark mode.
		                                        svgEl
		                                            .querySelectorAll('rect[data-owui-bg=\"1\"], rect.background, rect[class~=\"background\"], rect#background')
		                                            .forEach(removeNode);
		                                        try {{
		                                            const isWhiteish = (fill) => {{
		                                                const f = (fill || "").trim().toLowerCase();
	                                                return (
	                                                    f === "white" ||
	                                                    f === "#fff" ||
	                                                    f === "#ffffff" ||
	                                                    f === "rgb(255,255,255)" ||
	                                                    f === "rgb(255, 255, 255)"
	                                                );
	                                            }};
	                                            const nearly = (a, b) => Math.abs(a - b) <= 1e-3;
	                                            const rectMatches = (r, box) => {{
	                                                if (!box) return false;
	                                                const x = parseFloat(r.getAttribute("x") || "0");
	                                                const y = parseFloat(r.getAttribute("y") || "0");
	                                                const w = parseFloat(r.getAttribute("width") || "");
	                                                const h = parseFloat(r.getAttribute("height") || "");
	                                                if (!isFinite(x) || !isFinite(y) || !isFinite(w) || !isFinite(h)) return false;
	                                                return (
	                                                    nearly(x, box.minX) &&
	                                                    nearly(y, box.minY) &&
	                                                    nearly(w, box.width) &&
	                                                    nearly(h, box.height)
	                                                );
		                                            }};
		                                            const vbNow = parseViewBox(svgEl.getAttribute("viewBox"));
		                                            svgEl.querySelectorAll("rect[fill]").forEach((r) => {{
		                                                const fill = r.getAttribute("fill");
		                                                if (!isWhiteish(fill)) return;
		                                                if (rectMatches(r, vb0) || rectMatches(r, vbNow)) removeNode(r);
		                                            }});
		                                        }} catch (_e) {{}}
		                                        try {{
		                                            const vbCanvas = parseViewBox(svgEl.getAttribute(\"viewBox\")) || vb0 || vb;
		                                            if (vbCanvas) {{
		                                                const existing = svgEl.querySelector('rect[data-owui-canvas=\"1\"]');
		                                                const rect = existing || doc.createElementNS(\"http://www.w3.org/2000/svg\", \"rect\");
		                                                rect.setAttribute(\"data-owui-canvas\", \"1\");
		                                                rect.setAttribute(\"x\", `${{vbCanvas.minX}}`);
		                                                rect.setAttribute(\"y\", `${{vbCanvas.minY}}`);
		                                                rect.setAttribute(\"width\", `${{vbCanvas.width}}`);
		                                                rect.setAttribute(\"height\", `${{vbCanvas.height}}`);
		                                                rect.setAttribute(\"fill\", \"#FFFFFF\");
		                                                // Word quirk: without a full-canvas rect with *non-zero* opacity, Word will often
		                                                // only offer \"Convert to Shape\" when clicking on an actual stroke/fill (not empty space).
		                                                // We keep this rect nearly transparent and non-interactive.
		                                                rect.setAttribute(\"fill-opacity\", \"0.001\");
		                                                rect.setAttribute(\"stroke\", \"none\");
		                                                rect.setAttribute(\"stroke-opacity\", \"0\");
		                                                rect.setAttribute(\"pointer-events\", \"none\");
		                                                if (!existing) {{
		                                                    const first = svgEl.firstChild;
		                                                    svgEl.insertBefore(rect, first);
		                                                }}
		                                            }}
		                                        }} catch (_e) {{}}
		                                        return new XMLSerializer().serializeToString(svgEl);
		                                    }}
                                    function getMaxWidthEmu(xmlDoc) {{
                                        try {{
                                            const sects = xmlDoc.getElementsByTagName("w:sectPr");
                                            const sect = sects && sects.length ? sects[sects.length - 1] : null;
                                            if (!sect) return 5486400; // 6 in
                                            const pgSz = sect.getElementsByTagName("w:pgSz")[0];
                                            const pgMar = sect.getElementsByTagName("w:pgMar")[0];
                                            if (!pgSz || !pgMar) return 5486400;
                                            const pageW = parseInt(pgSz.getAttribute("w:w") || "", 10);
                                            const left = parseInt(pgMar.getAttribute("w:left") || "", 10);
                                            const right = parseInt(pgMar.getAttribute("w:right") || "", 10);
                                            if (!isFinite(pageW) || !isFinite(left) || !isFinite(right)) return 5486400;
                                            const twips = Math.max(1, pageW - left - right);
                                            return Math.round(twips * 635); // 1 twip = 635 EMU
                                        }} catch (_e) {{
                                            return 5486400;
                                        }}
                                    }}
                                    function getChildByTag(parent, tag) {{
                                        const nodes = parent.getElementsByTagName(tag);
                                        return nodes && nodes.length ? nodes[0] : null;
                                    }}
                                    try {{
                                        await loadScript(jszipUrl, "JSZip");
                                        await loadScript(mermaidUrl, "mermaid");
                                        // Mermaid init: disable htmlLabels to keep SVG Word-friendly; PNG fallback still included.
	                                        try {{
		                                            window.mermaid.initialize({{
		                                                startOnLoad: false,
		                                                theme: "default",
			                                                themeVariables: {{
			                                                    background: themeBackground,
			                                                    fontFamily: "Calibri, Segoe UI, Arial, sans-serif",
			                                                    fontSize: "10pt",
			                                                }},
			                                                themeCSS: ".slice {{ font-size: 10pt !important; }}\\n.legend text {{ font-size: 10pt !important; }}\\n.pieTitleText {{ font-size: 10pt !important; }}",
			                                                fontFamily: "Calibri, Segoe UI, Arial, sans-serif",
			                                                securityLevel: "strict",
			                                                flowchart: {{ htmlLabels: false }},
			                                            }});
	                                        }} catch (_e) {{
                                            // Ignore and proceed with defaults.
                                        }}
                                        const bytes = decodeBase64ToUint8Array(base64Data);
                                        const zip = new window.JSZip();
                                        await zip.loadAsync(bytes);
                                        const docXml = await zip.file("word/document.xml").async("string");
                                        const relsXml = await zip.file("word/_rels/document.xml.rels").async("string");
                                        const parser = new DOMParser();
                                        const xmlDoc = parser.parseFromString(docXml, "application/xml");
                                        const relsDoc = parser.parseFromString(relsXml, "application/xml");
                                        // Build rId -> target path mapping
                                        const rels = relsDoc.getElementsByTagName("Relationship");
                                        const rIdToTarget = {{}};
                                        for (let i = 0; i < rels.length; i++) {{
                                            const rel = rels[i];
                                            const id = rel.getAttribute("Id");
                                            const target = rel.getAttribute("Target");
                                            if (id && target) rIdToTarget[id] = target;
                                        }}
                                        const maxWidthEmu = getMaxWidthEmu(xmlDoc);
                                        const maxWidthEmuScaled = Math.max(1, Math.round(maxWidthEmu * Math.min(1.0, Math.max(0.1, displayScale || 1.0))));
                                        const drawings = xmlDoc.getElementsByTagName("w:drawing");
                                        const placeholders = [];
                                        for (let i = 0; i < drawings.length; i++) {{
                                            const drawing = drawings[i];
                                            const docPr = getChildByTag(drawing, "wp:docPr");
                                            if (!docPr) continue;
                                            const descr = docPr.getAttribute("descr") || "";
                                            if (!descr.startsWith("MERMAID_SRC:")) continue;
                                            const encoded = descr.substring("MERMAID_SRC:".length);
                                            const code = decodeURIComponent(encoded);
                                            const blip = getChildByTag(drawing, "a:blip");
                                            const ridPng = blip ? blip.getAttribute("r:embed") : null;
                                            const svgBlip = getChildByTag(drawing, "asvg:svgBlip");
                                            const ridSvg = svgBlip ? svgBlip.getAttribute("r:embed") : null;
                                            const container = getChildByTag(drawing, "wp:inline") || getChildByTag(drawing, "wp:anchor");
                                            const extent = container ? getChildByTag(container, "wp:extent") : null;
                                            const xfrm = getChildByTag(drawing, "a:xfrm");
                                            const xfrmExt = xfrm ? getChildByTag(xfrm, "a:ext") : null;
                                            placeholders.push({{ code, ridPng, ridSvg, extent, xfrmExt, svgBlip }});
                                        }}
                                        if (!placeholders.length) {{
                                            const blob = new Blob([bytes], {{ type: "application/vnd.openxmlformats-officedocument.wordprocessingml.document" }});
                                            downloadBlob(blob, filename);
                                            return;
                                        }}
                                        // Phase 1: Render all Mermaid diagrams sequentially (mermaid needs DOM)
                                        const renderResults = [];
                                        for (let i = 0; i < placeholders.length; i++) {{
                                            const item = placeholders[i];
                                            try {{
                                                const id = "owui-mermaid-" + i;
                                                const rendered = await window.mermaid.render(id, item.code);
                                                let svgText = rendered && rendered.svg ? rendered.svg : rendered;
                                                if (!svgText || typeof svgText !== "string") throw new Error("Mermaid returned empty SVG");
                                                svgText = normalizeSvgForWord(svgText);
                                                const hasForeignObject = /<foreignObject\\b/i.test(svgText);
                                                if (hasForeignObject && item.svgBlip) {{
                                                    try {{ item.svgBlip.parentNode && item.svgBlip.parentNode.removeChild(item.svgBlip); }} catch (_e) {{}}
                                                    item.ridSvg = null;
                                                }}
                                                const svgDoc = new DOMParser().parseFromString(svgText, "image/svg+xml");
                                                const svgEl = svgDoc.documentElement;
                                                const vb = parseViewBox(svgEl && svgEl.getAttribute ? svgEl.getAttribute("viewBox") : null);
                                                const ratio = vb && vb.width > 0 && vb.height > 0 ? (vb.width / vb.height) : (4/3);
                                                const widthEmu = maxWidthEmuScaled;
                                                const heightEmu = Math.max(1, Math.round(widthEmu / ratio));
                                                renderResults.push({{ item, svgText, widthEmu, heightEmu, success: true }});
                                            }} catch (err) {{
                                                console.error("Mermaid render failed for block", i, err);
                                                renderResults.push({{ item, svgText: null, widthEmu: 0, heightEmu: 0, success: false }});
                                            }}
                                        }}
                                        // Phase 2: Convert SVG to PNG in parallel for performance
                                        async function svgToPng(svgText, targetWidthPx, targetHeightPx) {{
                                            const canvas = document.createElement("canvas");
                                            const ctx = canvas.getContext("2d");
                                            const scale = Math.max(1.0, pngScale || 1.0);
                                            canvas.width = Math.round(targetWidthPx * scale);
                                            canvas.height = Math.round(targetHeightPx * scale);
                                            ctx.setTransform(1, 0, 0, 1, 0, 0);
                                            if (bgFill) {{
                                                ctx.fillStyle = bgFill;
                                                ctx.fillRect(0, 0, canvas.width, canvas.height);
                                            }}
                                            ctx.scale(scale, scale);
                                            const img = new Image();
                                            await new Promise((resolve, reject) => {{
                                                img.onload = resolve;
                                                img.onerror = reject;
                                                img.src = "data:image/svg+xml;base64," + btoa(unescape(encodeURIComponent(svgText)));
                                            }});
                                            ctx.drawImage(img, 0, 0, targetWidthPx, targetHeightPx);
                                            const pngDataUrl = canvas.toDataURL("image/png");
                                            return pngDataUrl.split(",")[1];
                                        }}
                                        // Create PNG conversion promises for parallel execution
                                        const pngPromises = renderResults.map(async (result, i) => {{
                                            if (!result.success || !result.svgText) return null;
                                            const {{ item, widthEmu, heightEmu }} = result;
                                            if (!item.ridPng || !rIdToTarget[item.ridPng]) return null;
                                            
                                            const targetWidthPx = Math.max(1, Math.round(widthEmu / 9525));
                                            const targetHeightPx = Math.max(1, Math.round(heightEmu / 9525));
                                            
                                            try {{
                                                const pngBase64 = await svgToPng(result.svgText, targetWidthPx, targetHeightPx);
                                                return {{ index: i, pngBase64, path: "word/" + rIdToTarget[item.ridPng] }};
                                            }} catch (err) {{
                                                console.error("PNG conversion failed for block", i, err);
                                                return null;
                                            }}
                                        }});
                                        // Wait for all PNG conversions to complete
                                        const pngResults = await Promise.all(pngPromises);
                                        // Phase 3: Update ZIP with all results
                                        for (let i = 0; i < renderResults.length; i++) {{
                                            const result = renderResults[i];
                                            if (!result.success) continue;
                                            
                                            const {{ item, svgText, widthEmu, heightEmu }} = result;
                                            
                                            // Update extent in XML
                                            if (item.extent) {{
                                                item.extent.setAttribute("cx", `${{widthEmu}}`);
                                                item.extent.setAttribute("cy", `${{heightEmu}}`);
                                            }}
                                            if (item.xfrmExt) {{
                                                item.xfrmExt.setAttribute("cx", `${{widthEmu}}`);
                                                item.xfrmExt.setAttribute("cy", `${{heightEmu}}`);
                                            }}
                                            // Write SVG part
                                            if (item.ridSvg && rIdToTarget[item.ridSvg]) {{
                                                zip.file("word/" + rIdToTarget[item.ridSvg], svgText);
                                            }}
                                        }}
                                        // Write PNG files from parallel results
                                        for (const pngResult of pngResults) {{
                                            if (pngResult && pngResult.pngBase64) {{
                                                zip.file(pngResult.path, pngResult.pngBase64, {{ base64: true }});
                                            }}
                                        }}
                                        const newDocXml = new XMLSerializer().serializeToString(xmlDoc);
                                        zip.file("word/document.xml", newDocXml);
                                        const finalBlob = await zip.generateAsync({{
                                            type: "blob",
                                            compression: "DEFLATE",
                                            compressionOptions: {{ level: 6 }},
                                        }});
                                        downloadBlob(finalBlob, filename);
                                    }} catch (error) {{
                                        console.error("Export pipeline failed:", error);
                                        const bytes = decodeBase64ToUint8Array(base64Data);
                                        const blob = new Blob([bytes], {{ type: "application/vnd.openxmlformats-officedocument.wordprocessingml.document" }});
                                        downloadBlob(blob, filename);
                                    }}
                                }})();
                                """
                            },
                        }
                    )
                await __event_emitter__(
                    {
                        "type": "status",
                        "data": {
                            "description": self._get_msg("exported"),
                            "done": True,
                        },
                    }
                )
                await self._emit_notification(
                    __event_emitter__,
                    self._get_msg("success", filename=filename),
                    "success",
                )
                return {"message": "Download triggered"}
            except Exception as e:
                logger.exception(f"Error exporting to Word: {str(e)}")
                await __event_emitter__(
                    {
                        "type": "status",
                        "data": {
                            "description": self._get_msg("export_failed", error=str(e)),
                            "done": True,
                        },
                    }
                )
                await self._emit_notification(
                    __event_emitter__,
                    self._get_msg("error_export", error=str(e)),
                    "error",
                )
    async def generate_title_using_ai(
        self, body: dict, content: str, user_id: str, request: Any
    ) -> str:
        if not request:
            return ""
        try:
            user_obj = Users.get_user_by_id(user_id)
            model = body.get("model")
            payload = {
                "model": model,
                "messages": [
                    {
                        "role": "system",
                        "content": "You are a helpful assistant. Generate a short, concise title (max 10 words) for the following text. Do not use quotes. Only output the title.",
                    },
                    {"role": "user", "content": content[:2000]},  # Limit content length
                ],
                "stream": False,
            }
            response = await generate_chat_completion(request, payload, user_obj)
            if response and "choices" in response:
                return response["choices"][0]["message"]["content"].strip()
        except Exception as e:
            logger.error(f"Error generating title: {e}")
        return ""
    def extract_title(self, content: str) -> str:
        """Extract title from Markdown h1/h2 only"""
        lines = content.split("\n")
        for line in lines:
            # Match h1-h2 headings only
            match = re.match(r"^#{1,2}\s+(.+)$", line.strip())
            if match:
                return match.group(1).strip()
        return ""
    def extract_chat_title(self, body: dict) -> str:
        """Extract chat title from common payload fields."""
        if not isinstance(body, dict):
            return ""
        candidates = []
        for key in ("chat", "conversation"):
            if isinstance(body.get(key), dict):
                candidates.append(body.get(key, {}).get("title", ""))
        for key in ("title", "chat_title"):
            value = body.get(key)
            if isinstance(value, str):
                candidates.append(value)
        for candidate in candidates:
            if candidate and isinstance(candidate, str):
                return candidate.strip()
        return ""
    def extract_chat_id(self, body: dict, metadata: Optional[dict]) -> str:
        """Extract chat_id from body or metadata"""
        if isinstance(body, dict):
            chat_id = body.get("chat_id") or body.get("id")
            if isinstance(chat_id, str) and chat_id.strip():
                return chat_id.strip()
            for key in ("chat", "conversation"):
                nested = body.get(key)
                if isinstance(nested, dict):
                    nested_id = nested.get("id") or nested.get("chat_id")
                    if isinstance(nested_id, str) and nested_id.strip():
                        return nested_id.strip()
        if isinstance(metadata, dict):
            chat_id = metadata.get("chat_id")
            if isinstance(chat_id, str) and chat_id.strip():
                return chat_id.strip()
        return ""
    async def fetch_chat_title(self, chat_id: str, user_id: str = "") -> str:
        """Fetch chat title from database by chat_id"""
        if not chat_id:
            return ""
        def _load_chat():
            if user_id:
                chat = Chats.get_chat_by_id_and_user_id(id=chat_id, user_id=user_id)
                if chat:
                    return chat
            return Chats.get_chat_by_id(chat_id)
        try:
            chat = await asyncio.to_thread(_load_chat)
        except Exception as exc:
            logger.warning(f"Failed to load chat {chat_id}: {exc}")
            return ""
        if not chat:
            return ""
        data = getattr(chat, "chat", {}) or {}
        title = data.get("title") or getattr(chat, "title", "")
        return title.strip() if isinstance(title, str) else ""
    def clean_filename(self, name: str) -> str:
        """Clean illegal characters from filename and strip emoji."""
        if not isinstance(name, str):
            return ""
        def _is_emoji_codepoint(codepoint: int) -> bool:
            # Common emoji ranges + flag regional indicators.
            return (
                0x1F000 <= codepoint <= 0x1FAFF
                or 0x1F1E6 <= codepoint <= 0x1F1FF
                or 0x2600 <= codepoint <= 0x26FF
                or 0x2700 <= codepoint <= 0x27BF
                or 0x2300 <= codepoint <= 0x23FF
                or 0x2B00 <= codepoint <= 0x2BFF
            )
        def _is_emoji_modifier(codepoint: int) -> bool:
            # VS15/VS16, ZWJ, keycap, skin tones, and tag characters used in some emoji sequences.
            return (
                codepoint in (0x200D, 0xFE0E, 0xFE0F, 0x20E3)
                or 0x1F3FB <= codepoint <= 0x1F3FF
                or 0xE0020 <= codepoint <= 0xE007F
            )
        without_emoji = "".join(
            ch
            for ch in name
            if not (_is_emoji_codepoint(ord(ch)) or _is_emoji_modifier(ord(ch)))
        )
        cleaned = re.sub(r'[\\/*?:"<>|]', "", without_emoji)
        cleaned = re.sub(r"\s+", " ", cleaned).strip().strip(".")
        return cleaned[:50].strip()
    def _max_embed_image_bytes(self) -> int:
        mb = getattr(self.valves, "MAX_EMBED_IMAGE_MB", 20)
        try:
            mb_i = int(mb)
        except Exception:
            mb_i = 20
        mb_i = max(1, mb_i)
        return mb_i * 1024 * 1024
    def _extract_owui_api_file_id(self, url: str) -> Optional[str]:
        if not isinstance(url, str) or not url:
            return None
        m = _OWUI_API_FILE_ID_RE.search(url)
        if not m:
            return None
        fid = (m.group("id") or "").strip()
        return fid or None
    def _read_file_bytes_limited(self, path: Path, max_bytes: int) -> Optional[bytes]:
        try:
            if not path.exists():
                return None
            try:
                size = path.stat().st_size
                if size > max_bytes:
                    return None
            except Exception:
                pass
            with path.open("rb") as f:
                data = f.read(max_bytes + 1)
            if len(data) > max_bytes:
                return None
            return data
        except Exception:
            return None
    def _decode_base64_limited(self, b64: str, max_bytes: int) -> Optional[bytes]:
        if not isinstance(b64, str):
            return None
        s = re.sub(r"\s+", "", b64.strip())
        if not s:
            return None
        # Rough pre-check: base64 expands by ~4/3. Avoid decoding clearly oversized payloads.
        est = (len(s) * 3) // 4
        if est > max_bytes:
            return None
        pad = (-len(s)) % 4
        if pad:
            s = s + ("=" * pad)
        try:
            out = base64.b64decode(s, validate=False)
        except (binascii.Error, ValueError):
            return None
        if len(out) > max_bytes:
            return None
        return out
    def _image_bytes_from_data_url(self, url: str, max_bytes: int) -> Optional[bytes]:
        if not isinstance(url, str):
            return None
        m = _DATA_IMAGE_URL_RE.match(url.strip())
        if not m:
            return None
        b64 = m.group("b64") or ""
        return self._decode_base64_limited(b64, max_bytes)
    def _read_from_s3(self, s3_path: str, max_bytes: int) -> Optional[bytes]:
        """Read file directly from S3 using environment variables for credentials."""
        if not BOTO3_AVAILABLE:
            return None
        # Parse s3://bucket/key
        if not s3_path.startswith("s3://"):
            return None
        path_without_prefix = s3_path[5:]  # Remove 's3://'
        parts = path_without_prefix.split("/", 1)
        if len(parts) < 2:
            return None
        bucket = parts[0]
        key = parts[1]
        # Read S3 config from environment variables
        endpoint_url = os.environ.get("S3_ENDPOINT_URL")
        access_key = os.environ.get("S3_ACCESS_KEY_ID")
        secret_key = os.environ.get("S3_SECRET_ACCESS_KEY")
        addressing_style = os.environ.get("S3_ADDRESSING_STYLE", "auto")
        if not all([endpoint_url, access_key, secret_key]):
            logger.debug(
                "S3 environment variables not fully configured, skipping S3 direct download."
            )
            return None
        try:
            s3_config = BotoConfig(
                s3={"addressing_style": addressing_style},
                connect_timeout=5,
                read_timeout=15,
            )
            s3_client = boto3.client(
                "s3",
                endpoint_url=endpoint_url,
                aws_access_key_id=access_key,
                aws_secret_access_key=secret_key,
                config=s3_config,
            )
            response = s3_client.get_object(Bucket=bucket, Key=key)
            body = response["Body"]
            data = body.read(max_bytes + 1)
            body.close()
            if len(data) > max_bytes:
                return None
            return data
        except Exception as e:
            logger.warning(f"S3 direct download failed for {s3_path}: {e}")
            return None
    def _image_bytes_from_owui_file_id(
        self, file_id: str, max_bytes: int
    ) -> Optional[bytes]:
        if not file_id:
            return None
        if Files is None:
            logger.error(
                "Files model is not available (import failed). Cannot retrieve file content."
            )
            return None
        try:
            file_obj = Files.get_file_by_id(file_id)
        except Exception as e:
            logger.error(f"Files.get_file_by_id({file_id}) failed: {e}")
            return None
        if not file_obj:
            logger.warning(f"File {file_id} not found in database.")
            return None
        # 1. Try data field (DB stored)
        data_field = getattr(file_obj, "data", None)
        if isinstance(data_field, dict):
            blob_value = data_field.get("bytes")
            if isinstance(blob_value, (bytes, bytearray)):
                raw = bytes(blob_value)
                return raw if len(raw) <= max_bytes else None
            for key in ("b64", "base64", "data"):
                inline = data_field.get(key)
                if isinstance(inline, str) and inline.strip():
                    return self._decode_base64_limited(inline, max_bytes)
        # 2. Try S3 direct download (fastest for object storage)
        s3_path = getattr(file_obj, "path", None)
        if isinstance(s3_path, str) and s3_path.startswith("s3://"):
            s3_data = self._read_from_s3(s3_path, max_bytes)
            if s3_data is not None:
                return s3_data
        # 3. Try file paths (Disk stored)
        # We try multiple path variations to be robust against CWD differences (e.g. Docker vs Local)
        for attr in ("path", "file_path", "absolute_path"):
            candidate = getattr(file_obj, attr, None)
            if isinstance(candidate, str) and candidate.strip():
                # Skip obviously non-local paths (S3, GCS, HTTP)
                if re.match(r"^(s3://|gs://|https?://)", candidate, re.IGNORECASE):
                    logger.debug(f"Skipping local read for non-local path: {candidate}")
                    continue
                p = Path(candidate)
                # Attempt 1: As-is (Absolute or relative to CWD)
                raw = self._read_file_bytes_limited(p, max_bytes)
                if raw is not None:
                    return raw
                # Attempt 2: Relative to ./data (Common in OpenWebUI)
                if not p.is_absolute():
                    try:
                        raw = self._read_file_bytes_limited(
                            Path("./data") / p, max_bytes
                        )
                        if raw is not None:
                            return raw
                    except Exception:
                        pass
                    # Attempt 3: Relative to /app/backend/data (Docker default)
                    try:
                        raw = self._read_file_bytes_limited(
                            Path("/app/backend/data") / p, max_bytes
                        )
                        if raw is not None:
                            return raw
                    except Exception:
                        pass
        # 4. Try URL (Object Storage / S3 Public URL)
        urls_to_try = []
        url_attr = getattr(file_obj, "url", None)
        if isinstance(url_attr, str) and url_attr:
            urls_to_try.append(url_attr)
        if isinstance(data_field, dict):
            url_data = data_field.get("url")
            if isinstance(url_data, str) and url_data:
                urls_to_try.append(url_data)
        if urls_to_try:
            import urllib.request
            for url in urls_to_try:
                if not url.startswith(("http://", "https://")):
                    continue
                try:
                    logger.info(
                        f"Attempting to download file {file_id} from URL: {url}"
                    )
                    # Use a timeout to avoid hanging
                    req = urllib.request.Request(
                        url, headers={"User-Agent": "OpenWebUI-Export-Plugin"}
                    )
                    with urllib.request.urlopen(req, timeout=15) as response:
                        if 200 <= response.status < 300:
                            data = response.read(max_bytes + 1)
                            if len(data) <= max_bytes:
                                return data
                            else:
                                logger.warning(
                                    f"File {file_id} from URL is too large (> {max_bytes} bytes)"
                                )
                except Exception as e:
                    logger.warning(f"Failed to download {file_id} from {url}: {e}")
        # 5. Try fetching via Local API (Last resort for S3/Object Storage without direct URL)
        # If we have the API token and base URL, we can try to fetch the content through the backend API.
        if self._api_base_url:
            api_url = f"{self._api_base_url}/api/v1/files/{file_id}/content"
            try:
                import urllib.request
                headers = {"User-Agent": "OpenWebUI-Export-Plugin"}
                if self._api_token:
                    headers["Authorization"] = self._api_token
                req = urllib.request.Request(api_url, headers=headers)
                with urllib.request.urlopen(req, timeout=15) as response:
                    if 200 <= response.status < 300:
                        data = response.read(max_bytes + 1)
                        if len(data) <= max_bytes:
                            return data
            except Exception:
                # API fetch failed, just fall through to the next method
                pass
        # 6. Try direct content attributes (last ditch)
        for attr in ("content", "blob", "data"):
            raw = getattr(file_obj, attr, None)
            if isinstance(raw, (bytes, bytearray)):
                b = bytes(raw)
                return b if len(b) <= max_bytes else None
        logger.warning(
            f"File {file_id} found but no content accessible. Attributes: {dir(file_obj)}"
        )
        return None
    def _add_image_placeholder(self, paragraph, alt: str, reason: str):
        label = (alt or "").strip() or "image"
        msg = f"[{label} not embedded: {reason}]"
        self._add_text_run(paragraph, msg, bold=False, italic=False, strike=False)
    def _try_embed_image(
        self, paragraph, image_bytes: bytes
    ) -> Tuple[bool, Optional[str]]:
        if not image_bytes:
            return False, "empty image bytes"
        try:
            run = paragraph.add_run()
            width = None
            if self._active_doc is not None:
                try:
                    width = self._available_block_width(self._active_doc)
                except Exception:
                    width = None
            run.add_picture(cast(Any, io.BytesIO(image_bytes)), width=width)
            return True, None
        except Exception as e:
            return False, str(e)
    def _embed_markdown_image(self, paragraph, alt: str, url: str):
        max_bytes = self._max_embed_image_bytes()
        u = (url or "").strip()
        if not u:
            self._add_image_placeholder(paragraph, alt, "missing URL")
            return
        image_bytes: Optional[bytes] = None
        if u.lower().startswith("data:"):
            image_bytes = self._image_bytes_from_data_url(u, max_bytes)
            if image_bytes is None:
                self._add_image_placeholder(
                    paragraph,
                    alt,
                    f"invalid data URL or exceeds {self.valves.MAX_EMBED_IMAGE_MB}MB",
                )
                return
        else:
            file_id = self._extract_owui_api_file_id(u)
            if not file_id:
                # External images are not fetched; treat as non-embeddable.
                self._add_image_placeholder(paragraph, alt, "external URL")
                return
            image_bytes = self._image_bytes_from_owui_file_id(file_id, max_bytes)
            if image_bytes is None:
                self._add_image_placeholder(
                    paragraph, alt, f"file unavailable ({file_id})"
                )
                return
        success, error_msg = self._try_embed_image(paragraph, image_bytes)
        if not success:
            self._add_image_placeholder(
                paragraph, alt, f"unsupported image type: {error_msg}"
            )
    async def markdown_to_docx(
        self,
        markdown_text: str,
        top_heading: str = "",
        has_h1: bool = False,
        sources: Optional[List[dict]] = None,
        event_emitter: Optional[Callable] = None,
    ) -> Document:
        """
        Convert Markdown text to Word document
        Supports: headings, paragraphs, bold, italic, code blocks, lists, tables, links
        Additionally: Mermaid fenced blocks (```mermaid) rendered client-side via Mermaid.js (SVG+PNG),
        LaTeX math to Word equations, and OpenWebUI citations to References.
        """
        doc = Document()
        self._active_doc = doc
        try:
            self._mermaid_figure_counter = 0
            self._mermaid_placeholder_counter = 0
            self._caption_style_name = None
            self._citation_anchor_by_index = {}
            self._citation_refs = self._build_citation_refs(sources or [])
            self._bookmark_id_counter = 1
            for ref in self._citation_refs:
                self._citation_anchor_by_index[ref.idx] = ref.anchor
            # Set default fonts
            self.set_document_default_font(doc)
            # If there is no h1 in content, prepend chat title as h1 when provided
            if top_heading and not has_h1:
                self.add_heading(doc, top_heading, 1)
            lines = markdown_text.split("\n")
            i = 0
            in_code_block = False
            code_block_content = []
            code_block_info_raw = ""
            code_block_lang = ""
            code_block_attrs: List[str] = []
            in_math_block = False
            math_block_delim = ""
            math_block_lines: List[str] = []
            in_list = False
            list_items = []
            list_type = None  # 'ordered' or 'unordered'
            total_lines = len(lines)
            last_update_time = time.time()
            while i < len(lines):
                # Update status every 2 seconds
                if event_emitter and time.time() - last_update_time > 2.0:
                    progress = int((i / total_lines) * 100)
                    await event_emitter(
                        {
                            "type": "status",
                            "data": {
                                "description": f"{self._get_msg('converting')} ({progress}%)",
                                "done": False,
                            },
                        }
                    )
                    last_update_time = time.time()
                line = lines[i]
                # Handle display math blocks (\[...\] or $$...$$)
                if not in_code_block and self.valves.MATH_ENABLE:
                    single_line = self._extract_single_line_math(line)
                    if single_line is not None:
                        if in_list and list_items:
                            self.add_list_to_doc(doc, list_items, list_type)
                            list_items = []
                            in_list = False
                        self._add_display_equation(doc, single_line)
                        i += 1
                        continue
                    if not in_math_block:
                        stripped = line.strip()
                        if stripped in (r"\[", "$$"):
                            if in_list and list_items:
                                self.add_list_to_doc(doc, list_items, list_type)
                                list_items = []
                                in_list = False
                            in_math_block = True
                            math_block_delim = stripped
                            math_block_lines = []
                            i += 1
                            continue
                    else:
                        stripped = line.strip()
                        close = r"\]" if math_block_delim == r"\[" else "$$"
                        if stripped == close:
                            in_math_block = False
                            latex = "\n".join(math_block_lines).strip()
                            self._add_display_equation(doc, latex)
                            math_block_delim = ""
                            math_block_lines = []
                            i += 1
                            continue
                        math_block_lines.append(line)
                        i += 1
                        continue
                # Handle code blocks
                if line.strip().startswith("```"):
                    if not in_code_block:
                        # Process pending list first
                        if in_list and list_items:
                            self.add_list_to_doc(doc, list_items, list_type)
                            list_items = []
                            in_list = False
                        in_code_block = True
                        code_block_info_raw = line.strip()[3:].strip()
                        code_block_lang, code_block_attrs = self._parse_fence_info(
                            code_block_info_raw
                        )
                        code_block_content = []
                    else:
                        # End code block
                        in_code_block = False
                        code_text = "\n".join(code_block_content)
                        if code_block_lang.lower() == "mermaid":
                            self._insert_mermaid_placeholder(doc, code_text)
                        else:
                            self.add_code_block(doc, code_text, code_block_lang)
                        code_block_content = []
                        code_block_info_raw = ""
                        code_block_lang = ""
                        code_block_attrs = []
                    i += 1
                    continue
                if in_code_block:
                    code_block_content.append(line)
                    i += 1
                    continue
                # Handle tables
                if line.strip().startswith("|") and line.strip().endswith("|"):
                    # Process pending list first
                    if in_list and list_items:
                        self.add_list_to_doc(doc, list_items, list_type)
                        list_items = []
                        in_list = False
                    table_lines = []
                    while i < len(lines) and lines[i].strip().startswith("|"):
                        table_lines.append(lines[i])
                        i += 1
                    self.add_table(doc, table_lines)
                    continue
                # Handle headings
                header_match = re.match(r"^(#{1,6})\s+(.+)$", line.strip())
                if header_match:
                    # Process pending list first
                    if in_list and list_items:
                        self.add_list_to_doc(doc, list_items, list_type)
                        list_items = []
                        in_list = False
                    level = len(header_match.group(1))
                    text = header_match.group(2)
                    self.add_heading(doc, text, level)
                    i += 1
                    continue
                # Handle unordered lists
                unordered_match = re.match(r"^(\s*)[-*+]\s+(.+)$", line)
                if unordered_match:
                    if not in_list or list_type != "unordered":
                        if in_list and list_items:
                            self.add_list_to_doc(doc, list_items, list_type)
                            list_items = []
                        in_list = True
                        list_type = "unordered"
                    indent = len(unordered_match.group(1)) // 2
                    list_items.append((indent, unordered_match.group(2)))
                    i += 1
                    continue
                # Handle ordered lists
                ordered_match = re.match(r"^(\s*)\d+[.)]\s+(.+)$", line)
                if ordered_match:
                    if not in_list or list_type != "ordered":
                        if in_list and list_items:
                            self.add_list_to_doc(doc, list_items, list_type)
                            list_items = []
                        in_list = True
                        list_type = "ordered"
                    indent = len(ordered_match.group(1)) // 2
                    list_items.append((indent, ordered_match.group(2)))
                    i += 1
                    continue
                # Handle blockquotes
                if line.strip().startswith(">"):
                    # Process pending list first
                    if in_list and list_items:
                        self.add_list_to_doc(doc, list_items, list_type)
                        list_items = []
                        in_list = False
                    # Collect consecutive quote lines
                    blockquote_lines = []
                    while i < len(lines) and lines[i].strip().startswith(">"):
                        # Remove leading > and optional space
                        quote_line = re.sub(r"^>\s?", "", lines[i])
                        blockquote_lines.append(quote_line)
                        i += 1
                    self.add_blockquote(doc, "\n".join(blockquote_lines))
                    continue
                # Handle horizontal rules
                if re.match(r"^[-*_]{3,}$", line.strip()):
                    # Process pending list first
                    if in_list and list_items:
                        self.add_list_to_doc(doc, list_items, list_type)
                        list_items = []
                        in_list = False
                    self.add_horizontal_rule(doc)
                    i += 1
                    continue
                # Handle empty lines
                if not line.strip():
                    # End list
                    if in_list and list_items:
                        self.add_list_to_doc(doc, list_items, list_type)
                        list_items = []
                        in_list = False
                    i += 1
                    continue
                # Handle normal paragraphs
                if in_list and list_items:
                    self.add_list_to_doc(doc, list_items, list_type)
                    list_items = []
                    in_list = False
                self.add_paragraph(doc, line)
                i += 1
            # Process remaining list
            if in_list and list_items:
                self.add_list_to_doc(doc, list_items, list_type)
            # If math block wasn't closed, render it as plain text for robustness.
            if in_math_block and math_block_lines:
                self.add_paragraph(doc, r"\[")
                for l in math_block_lines:
                    self.add_paragraph(doc, l)
                self.add_paragraph(doc, r"\]")
            if self._citation_refs:
                self._add_references_section(doc)
            return doc
        finally:
            self._active_doc = None
    def _extract_single_line_math(self, line: str) -> Optional[str]:
        s = line.strip()
        # \[ ... \]
        m = re.match(r"^\\\[(.*)\\\]$", s)
        if m:
            return m.group(1).strip()
        # $$ ... $$
        m = re.match(r"^\$\$(.*)\$\$$", s)
        if m:
            return m.group(1).strip()
        return None
    def _strip_reasoning_blocks(self, text: str) -> str:
        """
        Strip model reasoning blocks from assistant Markdown before export.
        OpenWebUI can include reasoning as interleaved <details type=\"reasoning\">...</details>
        (and sometimes <think>/<analysis> blocks). These should never be exported into DOCX.
        """
        if not text:
            return text
        cur = text
        for _ in range(10):
            prev = cur
            cur = _REASONING_DETAILS_RE.sub("", cur)
            cur = _THINK_RE.sub("", cur)
            cur = _ANALYSIS_RE.sub("", cur)
            if cur == prev:
                break
        # Clean up excessive blank lines left by removals.
        cur = re.sub(r"\n{4,}", "\n\n\n", cur)
        return cur
    def _add_display_equation(self, doc: Document, latex: str):
        latex = (latex or "").strip()
        if not latex:
            return
        if not LATEX_MATH_AVAILABLE:
            self.add_code_block(doc, latex, "latex")
            return
        try:
            mathml = latex_to_mathml(latex)
            omml = mathml2omml.convert(mathml)
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cast(Any, para)._p.append(self._wrap_omml_for_word(omml))
        except Exception as exc:
            logger.warning(f"Math conversion failed; falling back to text: {exc}")
            self.add_code_block(doc, latex, "latex")
    def _wrap_omml_for_word(self, omml: str):
        m_ns = "http://schemas.openxmlformats.org/officeDocument/2006/math"
        w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        # Keep the OMML payload as-is, but ensure it has the math namespace declared.
        xml = f'<m:oMathPara xmlns:m="{m_ns}" xmlns:w="{w_ns}">{omml}</m:oMathPara>'
        return parse_xml(xml)
    # (Math warning paragraphs removed)
    def _build_citation_refs(self, sources: List[dict]) -> List[_CitationRef]:
        citation_idx_map: Dict[str, int] = {}
        refs_by_idx: Dict[int, _CitationRef] = {}
        for source in sources or []:
            if not isinstance(source, dict):
                continue
            documents = source.get("document") or []
            metadatas = source.get("metadata") or []
            src_info = source.get("source") or {}
            src_name = src_info.get("name") if isinstance(src_info, dict) else None
            src_id_default = src_info.get("id") if isinstance(src_info, dict) else None
            src_urls = src_info.get("urls") if isinstance(src_info, dict) else None
            if not isinstance(documents, list):
                documents = []
            if not isinstance(metadatas, list):
                metadatas = []
            for idx_doc, _doc_text in enumerate(documents):
                meta = metadatas[idx_doc] if idx_doc < len(metadatas) else {}
                if not isinstance(meta, dict):
                    meta = {}
                source_id = meta.get("source") or src_id_default or "N/A"
                source_id_str = str(source_id)
                if source_id_str not in citation_idx_map:
                    citation_idx_map[source_id_str] = len(citation_idx_map) + 1
                idx = citation_idx_map[source_id_str]
                if idx in refs_by_idx:
                    continue
                url: Optional[str] = None
                if isinstance(source_id, str) and re.match(r"^https?://", source_id):
                    url = source_id
                elif isinstance(meta.get("url"), str) and re.match(
                    r"^https?://", meta["url"]
                ):
                    url = meta["url"]
                elif isinstance(src_urls, list) and src_urls:
                    if isinstance(src_urls[0], str) and re.match(
                        r"^https?://", src_urls[0]
                    ):
                        url = src_urls[0]
                title = (
                    (meta.get("title") if isinstance(meta.get("title"), str) else None)
                    or (meta.get("name") if isinstance(meta.get("name"), str) else None)
                    or (
                        src_name
                        if isinstance(src_name, str) and src_name.strip()
                        else None
                    )
                    or (url if url else None)
                    or source_id_str
                )
                anchor = f"OWUIRef{idx}"
                refs_by_idx[idx] = _CitationRef(
                    idx=idx,
                    anchor=anchor,
                    title=title,
                    url=url,
                    source_id=source_id_str,
                )
        return [refs_by_idx[i] for i in sorted(refs_by_idx.keys())]
    def _add_bookmark(self, paragraph, name: str):
        bookmark_id = self._bookmark_id_counter
        self._bookmark_id_counter += 1
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), str(bookmark_id))
        start.set(qn("w:name"), name)
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), str(bookmark_id))
        p = cast(Any, paragraph)._p
        p.insert(0, start)
        p.append(end)
    def _add_internal_hyperlink(self, paragraph, display_text: str, anchor: str):
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("w:anchor"), anchor)
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "Hyperlink")
        rPr.append(rStyle)
        new_run.append(rPr)
        t = OxmlElement("w:t")
        t.text = display_text
        new_run.append(t)
        hyperlink.append(new_run)
        cast(Any, paragraph)._p.append(hyperlink)
    def _add_references_section(self, doc: Document):
        self.add_heading(doc, self._get_msg("references"), 2)
        for ref in self._citation_refs:
            para = doc.add_paragraph(style="List Number")
            self._add_bookmark(para, ref.anchor)
            # Include URL as an external link when available.
            if ref.url:
                self._add_hyperlink(para, ref.title, ref.url, display_text=ref.title)
            else:
                self._add_text_run(
                    para, ref.title, bold=False, italic=False, strike=False
                )
    def _parse_fence_info(self, info_raw: str) -> Tuple[str, List[str]]:
        parts = [p for p in (info_raw or "").split() if p.strip()]
        if not parts:
            return "", []
        return parts[0], parts[1:]
    def _normalize_mermaid_text(self, source: str) -> str:
        text = (source or "").replace("\r\n", "\n").replace("\r", "\n").strip()
        return text + "\n"
    def _prepare_mermaid_for_js(self, source: str) -> str:
        """
        Prepare Mermaid source for client-side rendering:
        - strip title directives (caption already carries it),
        """
        text = self._strip_mermaid_title_for_render(source)
        return text
    def _png_with_text_chunk(self, png_bytes: bytes, keyword: str, value: str) -> bytes:
        """
        Ensure placeholder PNGs stay distinct in the DOCX package:
        python-docx may deduplicate identical image bytes into one media part.
        We insert a small tEXt chunk so each placeholder is unique, without changing
        dimensions or requiring external imaging libraries.
        """
        if not png_bytes.startswith(b"\x89PNG\r\n\x1a\n"):
            return png_bytes
        keyword_b = (keyword or "owui").encode("latin-1", errors="ignore")[:79]
        keyword_b = keyword_b.replace(b"\x00", b"") or b"owui"
        value_b = (value or "").encode("latin-1", errors="ignore")
        data = keyword_b + b"\x00" + value_b
        chunk_type = b"tEXt"
        crc = zlib.crc32(chunk_type + data) & 0xFFFFFFFF
        chunk = (
            struct.pack("!I", len(data)) + chunk_type + data + struct.pack("!I", crc)
        )
        out = bytearray()
        out.extend(png_bytes[:8])
        offset = 8
        inserted = False
        while offset + 8 <= len(png_bytes):
            length = struct.unpack("!I", png_bytes[offset : offset + 4])[0]
            ctype = png_bytes[offset + 4 : offset + 8]
            chunk_total = 12 + length
            if offset + chunk_total > len(png_bytes):
                break
            if ctype == b"IEND" and not inserted:
                out.extend(chunk)
                inserted = True
            out.extend(png_bytes[offset : offset + chunk_total])
            offset += chunk_total
            if ctype == b"IEND":
                break
        if not inserted:
            return png_bytes
        return bytes(out)
    def _make_mermaid_placeholder_png(self, seed: str) -> bytes:
        return self._png_with_text_chunk(_TRANSPARENT_1PX_PNG, "owui", seed)
    def _dummy_mermaid_svg_bytes(self) -> bytes:
        return (
            '<svg xmlns="http://www.w3.org/2000/svg" width="1" height="1" viewBox="0 0 1 1"></svg>'
        ).encode("utf-8")
    def _insert_mermaid_placeholder(self, doc: Document, mermaid_source: str):
        caption_title: Optional[str] = (
            self._extract_mermaid_title(mermaid_source)
            if self.valves.MERMAID_CAPTIONS_ENABLE
            else None
        )
        source_for_render = mermaid_source
        if self.valves.MERMAID_OPTIMIZE_LAYOUT:
            source_for_render = re.sub(
                r"^(graph|flowchart)\s+LR\b",
                r"\1 TD",
                source_for_render,
                flags=re.MULTILINE | re.IGNORECASE,
            )
        source_for_render = self._prepare_mermaid_for_js(source_for_render)
        self._mermaid_placeholder_counter += 1
        seed = hashlib.sha256(
            f"{self._mermaid_placeholder_counter}\n{source_for_render}".encode(
                "utf-8", errors="replace"
            )
        ).hexdigest()[:16]
        png_bytes = self._make_mermaid_placeholder_png(seed)
        try:
            shape = doc.add_picture(cast(Any, io.BytesIO(png_bytes)))
        except Exception as e:
            logger.warning(f"Failed to add Mermaid placeholder image: {e}")
            self.add_paragraph(doc, f"[Mermaid placeholder failed: {e}]")
            return
        try:
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass
        # Attach a dummy SVG part so we can later overwrite it client-side (SVG+PNG).
        self._attach_svg_blip(doc, shape, self._dummy_mermaid_svg_bytes())
        try:
            encoded = quote(source_for_render)
            inline = shape._inline
            docPr = inline.docPr
            docPr.set("descr", f"MERMAID_SRC:{encoded}")
            docPr.set("title", "Mermaid Diagram Placeholder")
        except Exception as exc:
            logger.warning(f"Failed to annotate Mermaid placeholder: {exc}")
        self._add_mermaid_caption(doc, caption_title)
    def _extract_mermaid_title(self, source: str) -> Optional[str]:
        lines = self._normalize_mermaid_text(source).split("\n")
        header_found = False
        for raw in lines:
            line = raw.strip()
            if not line:
                continue
            if line.startswith("%%{") and line.endswith("}%%"):
                continue
            if line.startswith("%%"):
                continue
            # diagram header line
            if not header_found:
                header_found = True
                # Mermaid beta/diagram headers can embed a title on the header line, e.g.:
                # - radar-beta title Foo
                # - xychart-beta title: "Foo"
                mt = re.match(
                    r"^(?P<header>\S.*?)(?:\s+title\s*:?\s+)(?P<title>.+)$",
                    line,
                    re.IGNORECASE,
                )
                if mt:
                    title = (mt.group("title") or "").strip().strip('"').strip("'")
                    if title:
                        return title
                continue
            # title "Foo" / title Foo
            m = re.match(r'^title\s*:?\s+"(.+)"\s*$', line, re.IGNORECASE)
            if m:
                return m.group(1).strip()
            m = re.match(r"^title\s*:?\s+(.+)$", line, re.IGNORECASE)
            if m:
                return m.group(1).strip().strip('"').strip("'")
        return None
    def _strip_mermaid_title_for_render(self, source: str) -> str:
        """
        Removes Mermaid title directives from the source before rendering.
        Captions already carry the title.
        """
        lines = self._normalize_mermaid_text(source).split("\n")
        out: List[str] = []
        header_found = False
        title_stripped = False
        meaningful_after_header = False
        for raw in lines:
            line = raw.rstrip("\n")
            stripped = line.strip()
            if not stripped:
                out.append(line)
                continue
            if stripped.startswith("%%{") and stripped.endswith("}%%"):
                out.append(line)
                continue
            if stripped.startswith("%%"):
                out.append(line)
                continue
            if not header_found:
                header_found = True
                # Some Mermaid diagram headers can embed a title on the header line, e.g.:
                # - radar-beta title Foo
                # - xychart-beta title: "Foo"
                mt = re.match(
                    r"^(?P<header>\S.*?)(?:\s+title\s*:?\s+)(?P<title>.+)$",
                    stripped,
                    re.IGNORECASE,
                )
                if mt:
                    cleaned = (mt.group("header") or "").strip()
                    out.append(cleaned if cleaned else stripped)
                    title_stripped = True
                    continue
                out.append(line)
                continue
            if not title_stripped and not meaningful_after_header:
                # Strip a standalone title directive line early in the diagram.
                if re.match(r'^title\s*:?\s+(".+"|.+)$', stripped, re.IGNORECASE):
                    title_stripped = True
                    continue
            # Consider this a meaningful content line after header.
            meaningful_after_header = True
            out.append(line)
        return "\n".join(out).strip() + "\n"
    def _ensure_caption_style(self, doc: Document) -> str:
        if self._caption_style_name is not None:
            return self._caption_style_name
        style_name = (self.valves.MERMAID_CAPTION_STYLE or "").strip()
        if style_name == "":
            # Empty means: do not apply a caption style.
            self._caption_style_name = ""
            return ""
        # Prefer existing style if present.
        try:
            _ = doc.styles[style_name]
            self._caption_style_name = style_name
            return style_name
        except KeyError:
            pass
        # If user requested "Caption" but it's missing, create a safe custom style name.
        if style_name.lower() == "caption":
            style_name = "OWUI Caption"
        try:
            _ = doc.styles[style_name]
            self._caption_style_name = style_name
            return style_name
        except KeyError:
            pass
        try:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = "Calibri"
            style.font.size = Pt(9)
            style.font.color.rgb = RGBColor(80, 80, 80)
            style.paragraph_format.space_before = Pt(2)
            style.paragraph_format.space_after = Pt(8)
            self._caption_style_name = style_name
            return style_name
        except Exception:
            self._caption_style_name = "Normal"
            return "Normal"
    def _add_mermaid_caption(self, doc: Document, title: Optional[str]):
        if not self.valves.MERMAID_CAPTIONS_ENABLE:
            return
        # Use configured prefix, or auto-detect from user language
        prefix = (self.valves.MERMAID_CAPTION_PREFIX or "").strip()
        if prefix == "":
            prefix = self._get_msg("figure_prefix")
        if prefix == "" and not title:
            return
        self._mermaid_figure_counter += 1
        if prefix == "":
            caption = title or ""
        else:
            base = f"{prefix} {self._mermaid_figure_counter}"
            caption = f"{base}: {title}" if title else base
        if caption == "":
            return
        para = doc.add_paragraph()
        style_name = self._ensure_caption_style(doc)
        if style_name:
            para.style = style_name
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.add_formatted_text(para, caption)
    def _available_block_width(self, doc: Document):
        section = doc.sections[0]
        return section.page_width - section.left_margin - section.right_margin
    def _attach_svg_blip(self, doc: Document, inline_shape: Any, svg_bytes: bytes):
        if not svg_bytes:
            return
        try:
            pkg = doc.part.package
            partname = pkg.next_partname("/word/media/image%d.svg")
            from docx.opc.part import Part
            svg_part = Part(partname, "image/svg+xml", svg_bytes)
            rid_svg = doc.part.relate_to(svg_part, RT.IMAGE)
            inline = inline_shape._inline
            blips = inline.xpath(".//a:blip")
            if not blips:
                return
            blip = blips[0]
            existing = blip.xpath(".//asvg:svgBlip")
            if existing:
                existing[0].set(qn("r:embed"), rid_svg)
                return
            extLst = OxmlElement("a:extLst")
            ext = OxmlElement("a:ext")
            ext.set("uri", "{96DAC541-7B7A-43D3-8B79-37D633B846F1}")
            svgBlip = OxmlElement("asvg:svgBlip")
            svgBlip.set(qn("r:embed"), rid_svg)
            ext.append(svgBlip)
            extLst.append(ext)
            blip.append(extLst)
        except Exception as exc:
            logger.warning(f"Failed to attach SVG blip; keeping PNG fallback: {exc}")
    # (Mermaid warning paragraphs removed)
    def set_document_default_font(self, doc: Document):
        """Set document default font using configured fonts."""
        style = doc.styles["Normal"]
        font = style.font
        font.name = self.valves.FONT_LATIN
        font.size = Pt(11)
        # Set Asian font
        style._element.rPr.rFonts.set(qn("w:eastAsia"), self.valves.FONT_ASIAN)
        # Set paragraph format
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph_format.space_after = Pt(6)
    def add_heading(self, doc: Document, text: str, level: int):
        """Add heading"""
        # Word heading levels start from 0, Markdown from 1
        heading_level = min(level, 9)  # Word supports up to Heading 9
        heading = doc.add_heading(level=heading_level)
        # Parse and add formatted text
        self.add_formatted_text(heading, text)
    def add_paragraph(self, doc: Document, text: str):
        """Add paragraph with inline formatting support"""
        paragraph = doc.add_paragraph()
        self.add_formatted_text(paragraph, text)
    def add_formatted_text(self, paragraph, text: str):
        """
        Parse Markdown inline formatting and add to paragraph.
        Supports: bold, italic, inline code, links, strikethrough, auto-link URLs,
        and inline LaTeX math \\(...\\) when MATH_ENABLE is on.
        """
        self._add_inline_segments(
            paragraph, text or "", bold=False, italic=False, strike=False
        )
    def _add_text_run(self, paragraph, s: str, bold: bool, italic: bool, strike: bool):
        if not s:
            return
        run = paragraph.add_run(s)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if strike:
            run.font.strike = True
    def _add_inline_code(self, paragraph, s: str):
        if s == "":
            return
        def _add_code_run(chunk: str):
            if not chunk:
                return
            run = paragraph.add_run(chunk)
            run.font.name = self.valves.FONT_CODE
            run._element.rPr.rFonts.set(qn("w:eastAsia"), self.valves.FONT_CODE)
            run.font.size = Pt(10)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "E8E8E8")
            run._element.rPr.append(shading)
        i = 0
        for m in _AUTO_URL_RE.finditer(s):
            start, end = m.span()
            if start > i:
                _add_code_run(s[i:start])
            raw = m.group(0)
            trimmed = raw
            while trimmed and trimmed[-1] in ".,;:!?)]}":
                trimmed = trimmed[:-1]
            suffix = raw[len(trimmed) :]
            normalized = self._normalize_url(trimmed)
            if normalized:
                self._add_hyperlink_code(
                    paragraph, display_text=trimmed, url=normalized
                )
            else:
                _add_code_run(raw)
            if suffix:
                _add_code_run(suffix)
            i = end
        if i < len(s):
            _add_code_run(s[i:])
    def _add_hyperlink_code(self, paragraph, display_text: str, url: str):
        u = self._normalize_url(url)
        if not u:
            self._add_inline_code(paragraph, display_text)
            return
        part = getattr(paragraph, "part", None)
        if part is None or not hasattr(part, "relate_to"):
            self._add_inline_code(paragraph, display_text)
            return
        r_id = part.relate_to(u, RT.HYPERLINK, is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), self.valves.FONT_CODE)
        rFonts.set(qn("w:hAnsi"), self.valves.FONT_CODE)
        rFonts.set(qn("w:eastAsia"), self.valves.FONT_CODE)
        rPr.append(rFonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "20")  # 10pt
        rPr.append(sz)
        sz_cs = OxmlElement("w:szCs")
        sz_cs.set(qn("w:val"), "20")
        rPr.append(sz_cs)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "E8E8E8")
        rPr.append(shading)
        new_run.append(rPr)
        t = OxmlElement("w:t")
        t.text = display_text
        new_run.append(t)
        hyperlink.append(new_run)
        cast(Any, paragraph)._p.append(hyperlink)
    def _add_inline_segments(
        self, paragraph, text: str, bold: bool, italic: bool, strike: bool
    ):
        i = 0
        n = len(text)
        def next_special(start: int) -> int:
            candidates = []
            for ch in ("`", "!", "[", "*", "_", "~", "$", "\\"):
                idx = text.find(ch, start)
                if idx != -1:
                    candidates.append(idx)
            idx = text.find(r"\(", start)
            if idx != -1:
                candidates.append(idx)
            idx = text.find("http://", start)
            if idx != -1:
                candidates.append(idx)
            idx = text.find("https://", start)
            if idx != -1:
                candidates.append(idx)
            idx = text.find("www.", start)
            if idx != -1:
                candidates.append(idx)
            return min(candidates) if candidates else n
        while i < n:
            # Markdown image: ![alt](url)
            if text.startswith("![", i):
                close = text.find("]", i + 2)
                if close != -1 and close + 1 < n and text[close + 1] == "(":
                    close_paren = text.find(")", close + 2)
                    if close_paren != -1:
                        alt = text[i + 2 : close]
                        url = text[close + 2 : close_paren].strip()
                        # Allow angle-bracket wrapped URLs: ![](</api/...>)
                        if url.startswith("<") and url.endswith(">") and len(url) >= 2:
                            url = url[1:-1].strip()
                        self._embed_markdown_image(paragraph, alt=alt, url=url)
                        i = close_paren + 1
                        continue
            if text[i] == "`":
                j = text.find("`", i + 1)
                if j != -1:
                    self._add_inline_code(paragraph, text[i + 1 : j])
                    i = j + 1
                    continue
            if text.startswith(r"\(", i):
                j = text.find(r"\)", i + 2)
                if j != -1:
                    self._add_inline_equation(
                        paragraph,
                        text[i + 2 : j],
                        bold=bold,
                        italic=italic,
                        strike=strike,
                    )
                    i = j + 2
                    continue
            # Handle backslash escapes
            if text[i] == "\\":
                if i + 1 < n:
                    ch = text[i + 1]
                    # Standard Markdown escapes + $ for math
                    if ch in "\\`*_{}[]()#+-.!|$":
                        self._add_text_run(paragraph, ch, bold, italic, strike)
                        i += 2
                        continue
                # Keep other backslashes literal
                self._add_text_run(paragraph, "\\", bold, italic, strike)
                i += 1
                continue
            # Handle long run of underscores (fill-in-the-blank)
            if text[i] == "_":
                run_len = 0
                while i + run_len < n and text[i + run_len] == "_":
                    run_len += 1
                if run_len >= 4:
                    self._add_text_run(
                        paragraph, text[i : i + run_len], bold, italic, strike
                    )
                    i += run_len
                    continue
            # Handle long run of asterisks (separator/mask)
            if text[i] == "*":
                run_len = 0
                while i + run_len < n and text[i + run_len] == "*":
                    run_len += 1
                if run_len >= 4:
                    self._add_text_run(
                        paragraph, text[i : i + run_len], bold, italic, strike
                    )
                    i += run_len
                    continue
            # Handle long run of tildes (separator)
            if text[i] == "~":
                run_len = 0
                while i + run_len < n and text[i + run_len] == "~":
                    run_len += 1
                if run_len >= 4:
                    self._add_text_run(
                        paragraph, text[i : i + run_len], bold, italic, strike
                    )
                    i += run_len
                    continue
            # Inline $...$ math (conservative parsing)
            if (
                text[i] == "$"
                and self.valves.MATH_ENABLE
                and self.valves.MATH_INLINE_DOLLAR_ENABLE
            ):
                # Avoid treating $$ as inline math here (block math uses $$ on its own line).
                if text.startswith("$$", i):
                    self._add_text_run(paragraph, "$", bold, italic, strike)
                    i += 1
                    continue
                # Markdown-ish heuristics to reduce false positives:
                # - Do not allow whitespace right after opening or right before closing
                # - Avoid cases like "USD$5" where opening is attached to an alnum
                if i + 1 >= n or text[i + 1].isspace():
                    self._add_text_run(paragraph, "$", bold, italic, strike)
                    i += 1
                    continue
                if i > 0 and text[i - 1].isalnum():
                    self._add_text_run(paragraph, "$", bold, italic, strike)
                    i += 1
                    continue
                j = i + 1
                while True:
                    j = text.find("$", j)
                    if j == -1:
                        break
                    # Skip escaped dollars inside: "\$"
                    if j > 0 and text[j - 1] == "\\":
                        j += 1
                        continue
                    break
                if j != -1:
                    inner = text[i + 1 : j]
                    if (
                        inner
                        and "\n" not in inner
                        and not inner[0].isspace()
                        and not inner[-1].isspace()
                    ):
                        # Treat "$5" as currency more often than math.
                        if _CURRENCY_NUMBER_RE.match(inner) and (
                            i == 0 or text[i - 1].isspace()
                        ):
                            self._add_text_run(paragraph, "$", bold, italic, strike)
                            i += 1
                            continue
                        # Disallow digit immediately following the closing $ (common in prices like "$5.00" already handled above).
                        if j + 1 < n and text[j + 1].isdigit():
                            self._add_text_run(paragraph, "$", bold, italic, strike)
                            i += 1
                            continue
                        self._add_inline_equation(
                            paragraph, inner, bold=bold, italic=italic, strike=strike
                        )
                        i = j + 1
                        continue
                self._add_text_run(paragraph, "$", bold, italic, strike)
                i += 1
                continue
            if text.startswith("~~", i):
                j = text.find("~~", i + 2)
                if j != -1:
                    self._add_inline_segments(
                        paragraph,
                        text[i + 2 : j],
                        bold=bold,
                        italic=italic,
                        strike=True,
                    )
                    i = j + 2
                    continue
            if text.startswith("**", i):
                j = text.find("**", i + 2)
                if j != -1:
                    self._add_inline_segments(
                        paragraph,
                        text[i + 2 : j],
                        bold=True,
                        italic=italic,
                        strike=strike,
                    )
                    i = j + 2
                    continue
            if text.startswith("__", i):
                j = text.find("__", i + 2)
                if j != -1:
                    self._add_inline_segments(
                        paragraph,
                        text[i + 2 : j],
                        bold=True,
                        italic=italic,
                        strike=strike,
                    )
                    i = j + 2
                    continue
            if text[i] == "*" and (i + 1 >= n or text[i + 1] != "*"):
                j = text.find("*", i + 1)
                if j != -1:
                    self._add_inline_segments(
                        paragraph,
                        text[i + 1 : j],
                        bold=bold,
                        italic=True,
                        strike=strike,
                    )
                    i = j + 1
                    continue
            if text[i] == "_" and (i + 1 >= n or text[i + 1] != "_"):
                j = text.find("_", i + 1)
                if j != -1:
                    self._add_inline_segments(
                        paragraph,
                        text[i + 1 : j],
                        bold=bold,
                        italic=True,
                        strike=strike,
                    )
                    i = j + 1
                    continue
            if text[i] == "[":
                close = text.find("]", i + 1)
                if close != -1 and close + 1 < n and text[close + 1] == "(":
                    close_paren = text.find(")", close + 2)
                    if close_paren != -1:
                        label = text[i + 1 : close]
                        url = text[close + 2 : close_paren]
                        self._add_hyperlink(paragraph, label, url)
                        i = close_paren + 1
                        continue
                # Citation marker like [12] -> internal link to References.
                if close != -1:
                    inner = text[i + 1 : close].strip()
                    if inner.isdigit():
                        idx = int(inner)
                        anchor = self._citation_anchor_by_index.get(idx)
                        if anchor:
                            self._add_internal_hyperlink(paragraph, f"[{idx}]", anchor)
                            i = close + 1
                            continue
            m = _AUTO_URL_RE.match(text, i)
            if m:
                raw = m.group(0)
                trimmed = raw
                while trimmed and trimmed[-1] in ".,;:!?)]}":
                    trimmed = trimmed[:-1]
                suffix = raw[len(trimmed) :]
                normalized = self._normalize_url(trimmed)
                if normalized:
                    # Display the original (trimmed) text; use normalized URL as the target.
                    self._add_hyperlink(
                        paragraph, trimmed, normalized, display_text=trimmed
                    )
                else:
                    self._add_text_run(paragraph, raw, bold, italic, strike)
                    i += len(raw)
                    continue
                if suffix:
                    self._add_text_run(paragraph, suffix, bold, italic, strike)
                i += len(raw)
                continue
            j = next_special(i)
            if j == i:
                # Unmatched special character; treat literally to avoid infinite loops.
                self._add_text_run(paragraph, text[i], bold, italic, strike)
                i += 1
            else:
                self._add_text_run(paragraph, text[i:j], bold, italic, strike)
                i = j
    def _normalize_url(self, url: str) -> str:
        u = (url or "").strip()
        if u.lower().startswith("www."):
            u = "https://" + u
        # Trim common trailing punctuation that often follows URLs in prose.
        while u and u[-1] in ".,;:!?)]}":
            u = u[:-1]
        return u
    def _add_hyperlink(
        self, paragraph, text: str, url: str, display_text: Optional[str] = None
    ):
        u = self._normalize_url(url)
        if not u:
            paragraph.add_run(display_text or text)
            return
        part = getattr(paragraph, "part", None)
        if part is None or not hasattr(part, "relate_to"):
            # Fallback if relationship API isn't available.
            run = paragraph.add_run(display_text or text)
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.font.underline = True
            return
        r_id = part.relate_to(u, RT.HYPERLINK, is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "Hyperlink")
        rPr.append(rStyle)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0000FF")
        rPr.append(color)
        u_el = OxmlElement("w:u")
        u_el.set(qn("w:val"), "single")
        rPr.append(u_el)
        new_run.append(rPr)
        t = OxmlElement("w:t")
        t.text = display_text or text
        new_run.append(t)
        hyperlink.append(new_run)
        cast(Any, paragraph)._p.append(hyperlink)
    def _add_inline_equation(
        self,
        paragraph,
        latex: str,
        bold: bool = False,
        italic: bool = False,
        strike: bool = False,
    ):
        latex = (latex or "").strip()
        if not latex:
            return
        if not self.valves.MATH_ENABLE or not LATEX_MATH_AVAILABLE:
            self._add_text_run(
                paragraph, f"\\({latex}\\)", bold=bold, italic=italic, strike=strike
            )
            return
        try:
            mathml = latex_to_mathml(latex)
            omml = mathml2omml.convert(mathml)
            o_math = self._omml_oMath_element(omml)
            run = paragraph.add_run()
            run.bold = bold
            run.italic = italic
            run.font.strike = strike
            cast(Any, run)._r.append(o_math)
        except Exception as exc:
            logger.warning(f"Inline math conversion failed; keeping literal: {exc}")
            self._add_text_run(
                paragraph, f"\\({latex}\\)", bold=bold, italic=italic, strike=strike
            )
    def _omml_oMath_element(self, omml: str):
        # Ensure the OMML element declares the math namespace so parse_xml works.
        m_ns = "http://schemas.openxmlformats.org/officeDocument/2006/math"
        s = (omml or "").strip()
        if s.startswith("<m:oMath>") and s.endswith("</m:oMath>"):
            inner = s[len("<m:oMath>") : -len("</m:oMath>")]
            s = f'<m:oMath xmlns:m="{m_ns}">{inner}</m:oMath>'
        elif s.startswith("<m:oMath") and "xmlns:m=" not in s.split(">", 1)[0]:
            s = s.replace("<m:oMath", f'<m:oMath xmlns:m="{m_ns}"', 1)
        return parse_xml(s)
    def add_code_block(self, doc: Document, code: str, language: str = ""):
        """Add code block with syntax highlighting"""
        # Token color mapping (based on common IDE themes)
        TOKEN_COLORS = {
            Token.Keyword: RGBColor(0, 92, 197),  # macOS blue - keywords
            Token.Keyword.Constant: RGBColor(0, 92, 197),
            Token.Keyword.Declaration: RGBColor(0, 92, 197),
            Token.Keyword.Namespace: RGBColor(0, 92, 197),
            Token.Keyword.Type: RGBColor(0, 92, 197),
            Token.Name.Function: RGBColor(0, 0, 0),  # Functions stay black
            Token.Name.Class: RGBColor(38, 82, 120),  # Deep cyan-blue - classes
            Token.Name.Decorator: RGBColor(170, 51, 0),  # Warm orange - decorators
            Token.Name.Builtin: RGBColor(0, 110, 71),  # Deep green - builtins
            Token.String: RGBColor(196, 26, 22),  # Red - strings
            Token.String.Doc: RGBColor(109, 120, 133),  # Gray - docstrings
            Token.Comment: RGBColor(109, 120, 133),  # Gray - comments
            Token.Comment.Single: RGBColor(109, 120, 133),
            Token.Comment.Multiline: RGBColor(109, 120, 133),
            Token.Number: RGBColor(28, 0, 207),  # Indigo - numbers
            Token.Number.Integer: RGBColor(28, 0, 207),
            Token.Number.Float: RGBColor(28, 0, 207),
            Token.Operator: RGBColor(90, 99, 120),  # Gray-blue - operators
            Token.Punctuation: RGBColor(0, 0, 0),  # Black - punctuation
        }
        def get_token_color(token_type):
            """Recursively find token color"""
            while token_type:
                if token_type in TOKEN_COLORS:
                    return TOKEN_COLORS[token_type]
                token_type = token_type.parent
            return None
        # Add language label if available
        if language:
            lang_para = doc.add_paragraph()
            lang_para.paragraph_format.space_before = Pt(6)
            lang_para.paragraph_format.space_after = Pt(0)
            lang_para.paragraph_format.left_indent = Cm(0.5)
            lang_run = lang_para.add_run(language.upper())
            lang_run.font.name = self.valves.FONT_CODE
            lang_run.font.size = Pt(8)
            lang_run.font.color.rgb = RGBColor(100, 100, 100)
            lang_run.font.bold = True
        # Add code block paragraph
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.5)
        paragraph.paragraph_format.space_before = Pt(3) if language else Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        # Add light gray background
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F7F7F7")
        paragraph._element.pPr.append(shading)
        # Try to use Pygments for syntax highlighting
        if PYGMENTS_AVAILABLE and language:
            try:
                lexer = get_lexer_by_name(language, stripall=False)
            except Exception:
                lexer = TextLexer()
            tokens = list(lex(code, lexer))
            for token_type, token_value in tokens:
                if not token_value:
                    continue
                run = paragraph.add_run(token_value)
                run.font.name = self.valves.FONT_CODE
                run._element.rPr.rFonts.set(qn("w:eastAsia"), self.valves.FONT_CODE)
                run.font.size = Pt(10)
                # Apply color
                color = get_token_color(token_type)
                if color:
                    run.font.color.rgb = color
                # Bold keywords
                if token_type in Token.Keyword:
                    run.font.bold = True
        else:
            # No syntax highlighting, plain text display
            run = paragraph.add_run(code)
            run.font.name = self.valves.FONT_CODE
            run._element.rPr.rFonts.set(qn("w:eastAsia"), self.valves.FONT_CODE)
            run.font.size = Pt(10)
    def add_table(self, doc: Document, table_lines: List[str]):
        """Add Markdown table with sane Word sizing/spacing, alignment, and hyperlinks/math support in cells."""
        if len(table_lines) < 2:
            return
        def _validate_hex(c: str, default: str) -> str:
            c = c.strip().lstrip("#")
            if re.fullmatch(r"[0-9A-Fa-f]{6}", c):
                return c
            return default
        header_fill = _validate_hex(self.valves.TABLE_HEADER_COLOR, "F2F2F2")
        zebra_fill = _validate_hex(self.valves.TABLE_ZEBRA_COLOR, "FBFBFB")
        def _split_row(line: str) -> List[str]:
            # Keep empty cells, trim surrounding pipes.
            raw = line.strip().strip("|")
            return [c.strip() for c in raw.split("|")]
        def _is_separator_row(cells: List[str]) -> bool:
            # Markdown separator: --- / :--- / ---: / :---:
            if not cells:
                return False
            ok = 0
            for c in cells:
                c = c.strip()
                if re.fullmatch(r":?-{3,}:?", c):
                    ok += 1
            return ok == len(cells)
        def _col_align(cell: str) -> WD_ALIGN_PARAGRAPH:
            s = (cell or "").strip()
            if s.startswith(":") and s.endswith(":"):
                return WD_ALIGN_PARAGRAPH.CENTER
            if s.endswith(":"):
                return WD_ALIGN_PARAGRAPH.RIGHT
            return WD_ALIGN_PARAGRAPH.LEFT
        def _set_cell_shading(cell, fill: str):
            tc_pr = cell._element.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), fill)
            tc_pr.append(shd)
        raw_rows = [_split_row(l) for l in table_lines if l.strip().startswith("|")]
        if not raw_rows:
            return
        sep_idx = 1 if len(raw_rows) > 1 and _is_separator_row(raw_rows[1]) else -1
        header = raw_rows[0]
        body = raw_rows[sep_idx + 1 :] if sep_idx >= 0 else raw_rows[1:]
        num_cols = max(len(header), *(len(r) for r in body)) if body else len(header)
        header = header + [""] * (num_cols - len(header))
        body = [r + [""] * (num_cols - len(r)) for r in body]
        aligns = [
            _col_align(c) for c in (raw_rows[1] if sep_idx == 1 else [""] * num_cols)
        ]
        table = doc.add_table(rows=1 + len(body), cols=num_cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        cast(Any, table).autofit = False
        # Cell margins (twips): smaller padding for compact tables.
        self._set_table_cell_margins(table, top=60, bottom=60, left=90, right=90)
        # Column widths: proportional to content, bounded, then normalized to page width.
        available_width = int(self._available_block_width(doc))
        min_col = max(int(Inches(0.55)), available_width // max(1, num_cols * 3))
        def _plain_len(s: str) -> int:
            t = re.sub(r"`([^`]+)`", r"\1", s or "")
            t = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", t)
            t = re.sub(r"\s+", " ", t).strip()
            return len(t)
        weights: List[int] = []
        for ci in range(num_cols):
            max_len = _plain_len(header[ci])
            for r in body:
                max_len = max(max_len, _plain_len(r[ci]))
            weights.append(max(1, min(max_len, 40)))
        sum_w = sum(weights) or 1
        widths = [max(min_col, int(available_width * w / sum_w)) for w in weights]
        total = sum(widths)
        if total > available_width:
            even = max(1, available_width // max(1, num_cols))
            widths = [even] * num_cols
            total = sum(widths)
        if total < available_width:
            rem = available_width - total
            order = sorted(range(num_cols), key=lambda i: weights[i], reverse=True)
            oi = 0
            while rem > 0 and order:
                widths[order[oi % len(order)]] += 1
                rem -= 1
                oi += 1
        for ci, w in enumerate(widths):
            table.columns[ci].width = w
            for row in table.rows:
                row.cells[ci].width = w
        def _format_cell_paragraph(para, align: WD_ALIGN_PARAGRAPH):
            para.alignment = align
            pf = para.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        def _fill_cell(cell, text: str, align: WD_ALIGN_PARAGRAPH, bold: bool = False):
            cell.text = ""
            parts = [
                p for p in re.split(r"(?:<br\s*/?>|\n)", text or "") if p is not None
            ]
            if not parts:
                parts = [""]
            for pi, part in enumerate(parts):
                para = cell.paragraphs[0] if pi == 0 else cell.add_paragraph()
                _format_cell_paragraph(para, align)
                self.add_formatted_text(para, part)
                for run in para.runs:
                    run.font.size = Pt(9)
                    if bold:
                        run.bold = True
        # Header row
        header_row = table.rows[0]
        self._set_table_header_row_repeat(header_row)
        for ci in range(num_cols):
            cell = header_row.cells[ci]
            _set_cell_shading(cell, header_fill)
            _fill_cell(
                cell,
                header[ci],
                aligns[ci] if ci < len(aligns) else WD_ALIGN_PARAGRAPH.LEFT,
                bold=True,
            )
        # Body rows
        for ri, row_data in enumerate(body, start=1):
            row = table.rows[ri]
            for ci in range(num_cols):
                cell = row.cells[ci]
                if (ri % 2) == 0:
                    _set_cell_shading(cell, zebra_fill)
                _fill_cell(
                    cell,
                    row_data[ci],
                    aligns[ci] if ci < len(aligns) else WD_ALIGN_PARAGRAPH.LEFT,
                )
    def _set_table_cell_margins(
        self, table, top: int, bottom: int, left: int, right: int
    ):
        tbl_pr = cast(Any, table)._tbl.tblPr
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        for tag, val in (
            ("top", top),
            ("bottom", bottom),
            ("left", left),
            ("right", right),
        ):
            el = OxmlElement(f"w:{tag}")
            el.set(qn("w:w"), str(int(val)))
            el.set(qn("w:type"), "dxa")
            tbl_cell_mar.append(el)
        tbl_pr.append(tbl_cell_mar)
    def _set_table_header_row_repeat(self, row):
        tr_pr = row._tr.get_or_add_trPr()
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)
    def add_list_to_doc(
        self, doc: Document, items: List[Tuple[int, str]], list_type: str
    ):
        """Add list"""
        for indent, text in items:
            paragraph = doc.add_paragraph()
            if list_type == "unordered":
                # Unordered list with bullets
                paragraph.style = "List Bullet"
            else:
                # Ordered list with numbers
                paragraph.style = "List Number"
            # Set indent
            paragraph.paragraph_format.left_indent = Cm(0.5 * (indent + 1))
            # Add formatted text
            self.add_formatted_text(paragraph, text)
    def add_horizontal_rule(self, doc: Document):
        """Add horizontal rule"""
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(12)
        # Add bottom border as horizontal rule
        pPr = paragraph._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        pBdr.append(bottom)
        pPr.append(pBdr)
    def add_blockquote(self, doc: Document, text: str):
        """Add blockquote with left border and gray background"""
        for line in text.split("\n"):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(1.0)
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(3)
            # Add left border
            pPr = paragraph._element.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "24")  # Border thickness
            left.set(qn("w:space"), "4")  # Space between border and text
            left.set(qn("w:color"), "CCCCCC")  # Gray border
            pBdr.append(left)
            pPr.append(pBdr)
            # Add light gray background
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "F9F9F9")
            pPr.append(shading)
            # Add formatted text
            self.add_formatted_text(paragraph, line)
            # Set font to italic gray
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(85, 85, 85)  # Dark gray text
                run.italic = True